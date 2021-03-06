# Field trial simulation
# guilherme borchardt - Oct., 2021
# simulation of inference about differences between averages.
# assumption: two independent variables that with similar variances and \
# follow normal distribution.
# inference statistics: T-Test for equal variance - two tails.

app_name = 'Field Trial Analyser'
app_version = 'v. 0.921' # code version

"""
this version (0.92):
    - add better DOCX file name
    - create folder /output to save DOCX files
    - save DOCX files into /output folder

    - delete *.pjg files
    - exit() bug on error function - to be corrected
"""

#print(chr(27) + "[2J")
print('loading....')

import winsound
import numpy as np
from prettytable import PrettyTable
from scipy import stats
import matplotlib.pyplot as plt
from datetime import date
import os
from statsmodels.stats import power as pwr
from art import *


    # from sys import *


    #
    # from FTA_data import product_segment_dict, product_list, specie_list, take_product
    #
    #
    # print(product_segment_dict['Gallipro'])
    # print(product_list)

def main():

    def value_p(p_input):
        if int(p_input) == 1:
            p_opt = 0.01
            return 0.01
        elif int(p_input) == 2:
            p_opt = 0.05
            return p_opt
        elif int(p_input) == 3:
            p_opt = 0.1
            return p_opt
        else:
            error()


    def error():
        # sound variables
        frequency = 2500  # Set Frequency To 2500 Hertz
        duration = 1000  # Set Duration To 1000 ms == 1 second

        art = text2art("ERROR", font='broadway')
        winsound.Beep(frequency, duration)
        print('\n')
        print(art)
        print('\nReview input data - program will be stopped!')
        os.system('pause')
        exit()


    # create folder for output files
    def create_folder(file):
        out_path = os.getcwd() + file

        if os.path.exists(out_path):
            return out_path
        else:
            os.mkdir(out_path)
            return out_path


    ############################################################################################

    out_path = create_folder('/output_files')  # create folder for files & return out_path

    cont = True
    while cont:
        os.system('cls') # clear console


        # build title - ASC II title
        art = text2art(f"Field trial analyser", font='smslant')
        print(art)  # print title

        ############################################################################################

        print('***** This APP is only suitable for continuous parameters that follow normal distribution *****\n')
        # define the input parameters
        print('Please input the data below:\n')
        trial_id = input('Trial ID: ').upper()


        product_name = input('Product Name: ').upper()


        def comma_finder_and_replacer(txt):
            txt.strip()  # remove eventual blank
            finder = txt.find(',')

            if finder != -1:  # case ',' exit
                txt = txt.replace(',', '.')

            return float(txt)


        c_mean = input('\nMean of control group: ')
        control_mean = comma_finder_and_replacer(c_mean)

        c_std_desv = input('Std. dev. of control group: ')
        control_std_desv = comma_finder_and_replacer(c_std_desv)

        t_mean = input(f'\nExpected Mean of {product_name} group: ')
        treatment_mean = comma_finder_and_replacer(t_mean)

        n_samples_per_treatment = int(input('\nNumber of repetitions / group: '))

        if n_samples_per_treatment <= 1:
            print('Num samples / group cannot be lower than 2')
            error()

        d_breakeven = input('\nDifference for break-even: ')
        delta_breakeven = comma_finder_and_replacer(d_breakeven)

        p_option = float(input('\nP value (type I error):\n[ 1 ] -> p=0,01\n[ 2 ] -> p=0,05\n[ 3 ] -> p=0,1\n'))
        # print()

        ############ internal variables ###############################################################################

        n_trials_repetitions = 5000  # number of repetitions to be simulated
        treatment_std_desv = control_std_desv  # consider standard deviation equal for both groups
        text_note = f'{app_name} - {app_version} - date: {date.today()} '
        n_lines = 15  # number of displayed trials on screen

        docx_file_name = f'Trial analyser ID = {trial_id} Product = {product_name} sample_size={n_samples_per_treatment}.docx'

        #############################################################################################

        title_font_size = 8
        color_no = '#ff6f69'
        color_yes = '#96ceb4'

        color_no2 = '#ffda6f'
        color_yes2 = '#7ed6c8'

        color_no3 = '#f14b23'
        color_yes3 = '#00b6bd'
        chart_dpi = 150

        # os.system('pause')

        # delta_mean = control_mean - treatment_mean

        delta_mean = treatment_mean - control_mean

        if delta_mean < 0:
            delta_negative = False
            # delta_breakeven = abs(delta_breakeven) * -1

        elif delta_mean == 0:
            print(f'Means from control and {product_name} group could not be equal!')
            error()

        elif delta_mean > 0:
            delta_negative = True
        # delta_breakeven = abs(delta_breakeven)

        p_level = value_p(p_option)

        print('-' * 80)

        # os.system('pause')

        # def main():

        x = PrettyTable()
        # print('\n@@@@ Field Trial Simulation @@@@\n')
        # print('User inputted values\n')
        # print(f'Control (no {product_name}) - {control_mean}')
        # print(f'Treatment (with {product_name}) - {treatment_mean}\n')
        # print(f'Std. desv. of groups: {control_std_desv}')
        # print(f'Number of repetitions / groups: {n_samples_per_treatment} ')
        #

        print(f'\nSimulated trial results from first {n_lines} simulated trials of {n_trials_repetitions} in total\n')
        x.field_names = ['run', 'n/treat.', 'Mean Control', f'Mean {product_name}', 'Delta', 'Break-even achieved?',
                         '(T-Test)', 'P', f'{product_name} Stat. approved?']
        count_expected = 0
        diff_expected = 0
        count_breakeven = 0

        for i in range(n_trials_repetitions):

            simul_data_control = np.random.normal(control_mean, control_std_desv, n_samples_per_treatment)
            simul_data_treatment = np.random.normal(treatment_mean, treatment_std_desv, n_samples_per_treatment)
            simul_control_average = round(np.average(simul_data_control), 3)
            simul_treatment_average = round(np.average(simul_data_treatment), 3)

            # modify ttest for one sided

            ttest, p = stats.ttest_ind(simul_data_control,
                                       simul_data_treatment,
                                       alternative='two-sided',
                                       )

            delta = round(simul_treatment_average - simul_control_average, 4)

            if p <= p_level:
                treat_accepted = 'YES'
                diff_expected = diff_expected + 1
            else:
                treat_accepted = '.'

            if delta_negative == True:  # when treatment reduce the mean
                if delta >= delta_mean:
                    expected = 'YES'
                    count_expected = count_expected + 1
                else:
                    expected = '.'

                if delta >= delta_breakeven:
                    breakeven = 'Yes'
                    count_breakeven = count_breakeven + 1
                else:
                    breakeven = '.'


            elif delta_negative == False:
                # delta = round(simul_control_average - simul_treatment_average, 4)

                if delta <= delta_mean:
                    expected = 'YES'
                    count_expected = count_expected + 1
                else:
                    expected = '.'

                if delta <= delta_breakeven:
                    breakeven = 'Yes'
                    count_breakeven = count_breakeven + 1
                else:
                    breakeven = '.'
                delta = abs(delta)
            else:
                print("verify your input")
                error()

            if i < n_lines:
                x.add_row([i + 1, n_samples_per_treatment, simul_control_average, simul_treatment_average, delta, breakeven,
                           round(ttest, 3), round(p, 3),
                           treat_accepted])

        print(str(x))

        # trial_customer_expected = [n_trials_repetitions - count_expected, count_expected]
        trial_breakeven_expected = [n_trials_repetitions - count_breakeven, count_breakeven]
        trial_statistic_diff_expected = [n_trials_repetitions - diff_expected, diff_expected]

        # percent_count_expected = round(count_expected/n_trials_repetitions * 100,1)
        percent_breakeven_expected = round(count_breakeven / n_trials_repetitions * 100, 1)
        percent_stat_expected = round(diff_expected / n_trials_repetitions * 100, 1)

        mean_delta = round(control_mean - treatment_mean, 3)

        # power calculation - review https://goodcalculators.com/effect-size-calculator/

        effect_s = abs((control_mean - treatment_mean) / control_std_desv)

        power_analysis = pwr.TTestIndPower()
        power_result = power_analysis.solve_power(
            effect_size=effect_s,  # effect size
            power=None,
            alpha=p_level,
            ratio=1,
            nobs1=n_samples_per_treatment,
        )
        type2_error = round((1 - power_result) * 100, 2)

        # sample size calculation - review https://goodcalculators.com/effect-size-calculator/
        sample_analysis = pwr.TTestIndPower()
        sample_result_2tail = sample_analysis.solve_power(
            effect_size=effect_s,
            power=.8,
            alpha=p_level,  # 2 sided
            ratio=1,
            nobs1=None,
        )
        sample_size_80_2tail = sample_result_2tail

        sample_result_1tail = sample_analysis.solve_power(
            effect_size=effect_s,
            power=.8,
            alpha=p_level * 2,  # 1 sided
            ratio=1,
            nobs1=None,
        )
        sample_size_80_1tail = sample_result_1tail

        print(f'\nTrial ID: {trial_id}  -  Product: {product_name}  -  Date: {date.today()}\n')
        print(
            f'\n- Trials showing high/lower break-even difference between Control and {product_name}: {percent_breakeven_expected} %')
        print(
            f'- Trials showing statistical significance between Control and {product_name} groups for P <= {p_level}: {percent_stat_expected} %\n')
        print('-' * 80)
        print('\n            Power & Sample Size (2-Sample t-Test)\n')
        print('-' * 80)
        print(f'\n- Effect size (Cohen??s d): {round(effect_s, 1)}')
        print(f'- Likelihood of Type II Error (false negative): {type2_error} %    -  ideal value is < 30%')
        print(f'- Statistical Power (1 - Type II Error): {round(100 - type2_error, 1)} %    -  ideal value is > 70%\n')
        print(
            f'- Sample size / group for Type II Error = 20% and Type I Error (P value) = {p_level}:\n   -{round(sample_size_80_2tail, 0)} - two-sided\n   -{round(sample_size_80_1tail, 0)} - one-sided\n')
        print('-' * 80)

        label = [f'NO', f'YES']
        if n_trials_repetitions >= 2:

            explode = [0, .05]
            plt.pie(trial_breakeven_expected,
                    labels=label,
                    startangle=90,
                    autopct='%1.1f%%',
                    colors=[color_no, color_yes],
                    explode=explode,
                    )
            plt.title(
                f'Likelihood of difference between groups equal/higher than break-even point ({delta_breakeven})\n(sample size = {n_samples_per_treatment})',
                fontsize=title_font_size,
                )
            # plt.legend(loc='best')
            axes = plt.gca()
            x_min, x_max = axes.get_xlim()
            y_min, y_max = axes.get_ylim()
            plt.text(x_min,
                     y_min - .15,
                     text_note,
                     fontsize=6,
                     alpha=.5
                     )

            plt.savefig(f'{out_path}/{product_name} trial_breakeven_n={n_samples_per_treatment}.jpg',
                        dpi=chart_dpi,
                        )
            plt.show()

            # plt.pie(trial_customer_expected,
            #         labels=label,
            #         startangle=90,
            #         autopct='%1.1f%%',
            #         colors=[color_no2,color_yes2],
            #         explode = explode,
            #         )
            # plt.title(f'Likelihood for difference higher/lower of {mean_delta} between groups?\n(sample size = {n_samples_per_treatment})',
            #           fontsize= title_font_size,
            #           )
            # #plt.legend(loc='best')
            # axes = plt.gca()
            # x_min, x_max = axes.get_xlim()
            # y_min, y_max = axes.get_ylim()
            # plt.text(x_min,
            #          y_min - .15,
            #          text_note,
            #          fontsize=6,
            #          alpha=.5
            #          )
            # plt.savefig(f'{product_name} trial_difference_n={n_samples_per_treatment}.jpg',
            #             dpi=chart_dpi,
            #             )
            # plt.show()
            #

            plt.pie(trial_statistic_diff_expected,
                    labels=label,
                    startangle=90,
                    autopct='%1.1f%%',
                    colors=[color_no3, color_yes3],
                    explode=explode,
                    )
            plt.title(
                f'Likelihood for statistic significance (P<={p_level}) between groups?\n(sample size = {n_samples_per_treatment})',
                fontsize=title_font_size,
                )
            # plt.legend(loc='best')
            axes = plt.gca()
            x_min, x_max = axes.get_xlim()
            y_min, y_max = axes.get_ylim()
            plt.text(x_min,
                     y_min - .15,
                     f'{text_note} - Type II error rate: {type2_error} %',
                     fontsize=6,
                     alpha=.5
                     )
            # plt.savefig(f'{product_name} trial_simulation_statistics_n={n_samples_per_treatment}.jpg',
            #             dpi=chart_dpi,
            #             )
            plt.savefig(f'{out_path}/{product_name} trial_simulation_statistics_n={n_samples_per_treatment}.jpg',
                        dpi=chart_dpi,
                        )

            plt.show()

            print(f'DOCX file with report was saved into "{os.getcwd()}\output_files"  folder\n')
            print(f'                                              Field trial analyser - {app_version}\n')
            print('-' * 80)

            # create word document

            import docx

            docx_lines = 70
            output_doc = docx.Document()
            # document title
            output_doc.add_heading(f'{app_name}', 0)
            # output_doc.add_paragraph(f'             {text_note}')

            # document body
            output_doc.add_paragraph('User input Data:')
            output_doc.add_paragraph(f'Trial ID: {trial_id} - Product Name: {product_name}  -  Date: {date.today()}')
            output_doc.add_paragraph(f'Sample size / group:  {n_samples_per_treatment}')
            output_doc.add_paragraph(f'Mean from Control group: {control_mean}  sd: {control_std_desv}')
            output_doc.add_paragraph(f'Mean from {product_name} group: {treatment_mean}  sd: {control_std_desv}')
            output_doc.add_paragraph(
                f'Delta of means: {round(delta_mean, 3)} ({round(delta_mean / control_mean * 100, 1)} %)')
            output_doc.add_paragraph(
                f'Min. acceptable delta: {delta_breakeven} ({round(delta_breakeven / control_mean * 100, 1)} %)')
            output_doc.add_paragraph(f'P value <= {p_level}')

            output_doc.add_paragraph('=' * docx_lines)

            output_doc.add_paragraph('Output:')
            output_doc.add_paragraph(
                f' - Likelihood of show an accepted delta lower/higher of {round(delta_breakeven, 3)}: {percent_breakeven_expected} %')
            output_doc.add_paragraph(
                f' - Likelihood of show statistical significance between Control and {product_name} groups: {percent_stat_expected} % \r')

            output_doc.add_paragraph('=' * docx_lines)

            output_doc.add_paragraph('Power Analysis (2-Sample t-Test)')
            output_doc.add_paragraph(
                f'- Likelihood of Type II Error (false negative): {type2_error} %  -  ideal value is <30%')
            output_doc.add_paragraph(
                f'- Statistical Power (1 - Type II Error): {round(100 - type2_error, 1)} %  -  ideal value is >70%')
            output_doc.add_paragraph('=' * docx_lines)
            output_doc.add_paragraph('Ideal Sample size')
            output_doc.add_paragraph(
                f' - Sample size / group for statistic significance:  (Power of 80% and P value = {p_level}):\r         -{round(sample_size_80_2tail, 0)} - two-sided\r         -{round(sample_size_80_1tail, 0)} - one-sided\r')
            output_doc.add_paragraph('=' * docx_lines)

            output_doc.add_page_break()

            output_doc.add_picture(f'{out_path}/{product_name} trial_breakeven_n={n_samples_per_treatment}.jpg')
            output_doc.add_paragraph('')
            output_doc.add_picture(f'{out_path}/{product_name} trial_simulation_statistics_n={n_samples_per_treatment}.jpg')
            output_doc.add_paragraph(f'App version: {app_version}')

            output_doc.save(f'{out_path}/{docx_file_name}')

            # delete jpg files created by matplotlib
            try:
                os.remove(f'{out_path}/{product_name} trial_simulation_statistics_n={n_samples_per_treatment}.jpg')
                os.remove(f'{out_path}/{product_name} trial_breakeven_n={n_samples_per_treatment}.jpg')
            except:
                print('cound not remove JPG files')
            else:
                print('JPG files removed')


            # import sys
            # sys.path.append(out_path)
            # import subprocess
            # subprocess.Popen(f'explorer"{out_path}"')

            def repeat(user_option):
                if user_option == 'Y' or user_option == 'y':
                    c = True
                else:
                    c = False
                return c


            cont_loop = input('Do You want to analyse other trial? (Y/N)\n')
            cont = repeat(cont_loop)  # if Y repeat programm

    #
#
if __name__ ==  '__main__':
    main()






